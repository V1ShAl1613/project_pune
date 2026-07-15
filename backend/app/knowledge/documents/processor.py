from __future__ import annotations

import csv
import io
import json
import re
import xml.etree.ElementTree as ElementTree
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from app.core.settings import AppSettings
from app.knowledge.validators import KnowledgeValidator


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        text = data.strip()
        if text:
            self.parts.append(text)

    def extract(self) -> str:
        return " ".join(self.parts)


@dataclass(slots=True)
class DocumentProcessor:
    """Extract text, metadata, and structural hints from supported document types."""

    settings: AppSettings
    validator: KnowledgeValidator

    def process(self, content: str, *, file_name: str, file_type: str, mime_type: str | None = None) -> dict[str, Any]:
        normalized_type = file_type.lower().lstrip(".")
        extractor = self._extractors().get(normalized_type, self._extract_plain_text)
        extracted = extractor(content, file_name=file_name, file_type=normalized_type)
        text = self._clean_text(extracted["text"])
        text = self._normalize_encoding(text)
        language = self.validator.language_guess(text)
        checksum = self.validator.checksum(text)
        return {
            "text": text,
            "metadata": {
                **extracted["metadata"],
                "file_name": file_name,
                "file_type": normalized_type,
                "mime_type": mime_type or self.validator.infer_mime_type(file_name, normalized_type),
                "language": language,
                "encoding": self.validator.detect_encoding(text),
                "checksum": checksum,
                "word_count": self._word_count(text),
                "char_count": len(text),
                "page_count": extracted["metadata"].get("page_count"),
            },
            "language": language,
            "checksum": checksum,
        }

    def read_file(self, path: str) -> tuple[str, str | None]:
        file_path = Path(path)
        file_type = file_path.suffix.lower().lstrip(".")
        if file_type == "pdf":
            return self._extract_pdf(file_path.read_bytes()), self.validator.infer_mime_type(file_path.name, file_type)
        if file_type == "docx":
            return self._extract_docx(file_path.read_bytes()), self.validator.infer_mime_type(file_path.name, file_type)
        return file_path.read_text(encoding="utf-8", errors="ignore"), self.validator.infer_mime_type(file_path.name, file_type)

    def _extractors(self):
        return {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_plain_text,
            "md": self._extract_markdown,
            "csv": self._extract_csv,
            "json": self._extract_json,
            "xml": self._extract_xml,
            "html": self._extract_html,
            "yaml": self._extract_yaml,
            "yml": self._extract_yaml,
            "log": self._extract_plain_text,
        }

    def _extract_plain_text(self, content: str, **_: Any) -> dict[str, Any]:
        return {"text": content, "metadata": {"detected_format": "text"}}

    def _extract_markdown(self, content: str, **_: Any) -> dict[str, Any]:
        headings = [line.lstrip("# ") for line in content.splitlines() if line.strip().startswith("#")]
        return {"text": content, "metadata": {"detected_format": "markdown", "headings": headings, "section_count": len(headings)}}

    def _extract_csv(self, content: str, **_: Any) -> dict[str, Any]:
        rows = list(csv.reader(io.StringIO(content)))
        header = rows[0] if rows else []
        return {"text": content, "metadata": {"detected_format": "csv", "row_count": len(rows), "column_count": len(header), "header": header}}

    def _extract_json(self, content: str, **_: Any) -> dict[str, Any]:
        parsed = json.loads(content or "{}")
        return {"text": json.dumps(parsed, ensure_ascii=False, indent=2), "metadata": {"detected_format": "json", "keys": list(parsed) if isinstance(parsed, dict) else []}}

    def _extract_xml(self, content: str, **_: Any) -> dict[str, Any]:
        root = ElementTree.fromstring(content)
        return {"text": content, "metadata": {"detected_format": "xml", "root_tag": root.tag, "element_count": len(list(root.iter()))}}

    def _extract_html(self, content: str, **_: Any) -> dict[str, Any]:
        extractor = _HTMLTextExtractor()
        extractor.feed(content)
        text = extractor.extract()
        return {"text": text, "metadata": {"detected_format": "html", "title": self._find_html_title(content), "section_count": content.lower().count("<section")}}

    def _extract_yaml(self, content: str, **_: Any) -> dict[str, Any]:
        return {"text": content, "metadata": {"detected_format": "yaml", "line_count": len(content.splitlines())}}

    def _extract_pdf(self, content: bytes | str, **_: Any) -> dict[str, Any]:
        if isinstance(content, str):
            content_bytes = content.encode("utf-8", errors="ignore")
        else:
            content_bytes = content
        text = self._try_pdf_extract(content_bytes)
        return {"text": text, "metadata": {"detected_format": "pdf", "page_count": self._pdf_page_count(content_bytes)}}

    def _extract_docx(self, content: bytes | str, **_: Any) -> dict[str, Any]:
        if isinstance(content, str):
            content_bytes = content.encode("utf-8", errors="ignore")
        else:
            content_bytes = content
        text = self._try_docx_extract(content_bytes)
        return {"text": text, "metadata": {"detected_format": "docx"}}

    def _try_pdf_extract(self, content: bytes) -> str:
        try:
            import fitz  # type: ignore

            document = fitz.open(stream=content, filetype="pdf")
            pages = [page.get_text("text") for page in document]
            return "\n".join(pages)
        except Exception:
            return content.decode("utf-8", errors="ignore")

    def _pdf_page_count(self, content: bytes) -> int | None:
        try:
            import fitz  # type: ignore

            document = fitz.open(stream=content, filetype="pdf")
            return document.page_count
        except Exception:
            return None

    def _try_docx_extract(self, content: bytes) -> str:
        try:
            from docx import Document  # type: ignore

            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            tables = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            return "\n".join([*paragraphs, *tables])
        except Exception:
            return content.decode("utf-8", errors="ignore")

    def _find_html_title(self, content: str) -> str | None:
        match = re.search(r"<title>(.*?)</title>", content, re.IGNORECASE | re.DOTALL)
        return match.group(1).strip() if match else None

    def _clean_text(self, text: str) -> str:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()

    def _normalize_encoding(self, text: str) -> str:
        if not self.settings.knowledge_encoding_validation_enabled:
            return text
        return text.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")

    def _word_count(self, text: str) -> int:
        return len([token for token in re.split(r"\s+", text.strip()) if token])
